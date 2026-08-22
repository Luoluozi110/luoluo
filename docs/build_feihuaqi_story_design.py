from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "飞花棋-剧情设计文档.docx"

# compact_reference_guide preset + named CJK/title/callout overrides
ASCII_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
INK = "1F2933"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "66717C"
GOLD = "A97920"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=None, bold=None, italic=None, color=None, ascii_font=ASCII_FONT, cjk_font=CJK_FONT):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), cjk_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:eastAsia"), "zh-CN")
    return run


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D6DCE3", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row):
    repeat_table_header(row)


def shade_paragraph(paragraph, fill, left_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if left_color:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "22")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_color)
        borders.append(left)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = ASCII_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = ASCII_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        srpr = style._element.get_or_add_rPr()
        srfonts = srpr.rFonts
        if srfonts is None:
            srfonts = OxmlElement("w:rFonts")
            srpr.insert(0, srfonts)
        srfonts.set(qn("w:ascii"), ASCII_FONT)
        srfonts.set(qn("w:hAnsi"), ASCII_FONT)
        srfonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    subtitle = styles["Subtitle"]
    subtitle.font.name = ASCII_FONT
    subtitle.font.size = Pt(14)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(12)
    srpr = subtitle._element.get_or_add_rPr()
    srfonts = srpr.rFonts
    if srfonts is None:
        srfonts = OxmlElement("w:rFonts")
        srpr.insert(0, srfonts)
    srfonts.set(qn("w:ascii"), ASCII_FONT)
    srfonts.set(qn("w:hAnsi"), ASCII_FONT)
    srfonts.set(qn("w:eastAsia"), CJK_FONT)

    code = styles.add_style("Story Code", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    crpr = code._element.get_or_add_rPr()
    crfonts = crpr.rFonts
    if crfonts is None:
        crfonts = OxmlElement("w:rFonts")
        crpr.insert(0, crfonts)
    crfonts.set(qn("w:ascii"), "Consolas")
    crfonts.set(qn("w:hAnsi"), "Consolas")
    crfonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    code.paragraph_format.left_indent = Inches(0.2)
    code.paragraph_format.right_indent = Inches(0.2)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.0


def add_numbering_definition(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abs_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abs_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), ASCII_FONT)
    r_fonts.set(qn("w:hAnsi"), ASCII_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    # OOXML requires every abstractNum to precede every num. Word may silently
    # reinterpret lists when these elements are interleaved.
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def add_para(doc, text="", *, bold_prefix=None, italic=False, color=None, align=None, after=None, keep=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=color or INK)
        set_run_font(p.add_run(text[len(bold_prefix):]), italic=italic, color=color or INK)
    else:
        set_run_font(p.add_run(text), italic=italic, color=color or INK)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if keep:
        p.paragraph_format.keep_with_next = True
    return p


def add_bullet(doc, text, bullet_num_id, bold_prefix=None):
    p = doc.add_paragraph()
    apply_numbering(p, bullet_num_id)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_numbered(doc, text, num_id, title=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    if title and text.startswith(title):
        set_run_font(p.add_run(title), bold=True, color=DARK_BLUE)
        set_run_font(p.add_run(text[len(title):]))
    else:
        set_run_font(p.add_run(text))
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_callout(doc, label, text, color=GOLD, fill=CALLOUT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, fill, color)
    set_run_font(p.add_run(f"{label}  "), bold=True, color=color)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_quote(doc, text, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.2
    shade_paragraph(p, "F8FAFC", BLUE)
    set_run_font(p.add_run(text), size=10.5, italic=True, color=NAVY)
    if attribution:
        set_run_font(p.add_run(f"\n——{attribution}"), size=9, color=MUTED)
    return p


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        p.paragraph_format.page_break_before = True
    set_run_font(p.add_run(text), size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_tagged(doc, tag, text, tag_color=BLUE):
    p = doc.add_paragraph()
    set_run_font(p.add_run(f"【{tag}】"), bold=True, color=tag_color)
    set_run_font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(header), size=9.5, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if idx == 0:
                set_run_font(p.add_run(str(value)), size=9.2, bold=True, color=DARK_BLUE)
            else:
                set_run_font(p.add_run(str(value)), size=9.2, color=INK)
        if len(table.rows) % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_character(doc, name, role, desire, belief, weakness, arc, voice, sample_lines, bullet_num_id):
    add_heading(doc, f"{name}｜{role}", 3)
    add_para(doc, f"欲望：{desire}", bold_prefix="欲望：")
    add_para(doc, f"信念：{belief}", bold_prefix="信念：")
    add_para(doc, f"裂缝：{weakness}", bold_prefix="裂缝：")
    add_para(doc, f"角色弧：{arc}", bold_prefix="角色弧：")
    add_para(doc, f"声音：{voice}", bold_prefix="声音：")
    for line in sample_lines:
        add_bullet(doc, line, bullet_num_id)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("飞花棋｜剧情设计文档"), size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("内部创作参考 · 2026-08-22 · 第 "), size=8.5, color=MUTED)
    add_page_field(fp)
    set_run_font(fp.add_run(" 页"), size=8.5, color=MUTED)

    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0]
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(ffp.add_run("基于 2026-08-22 工作区文本整理｜建议稿"), size=8.5, color=MUTED)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build_document():
    doc = Document()
    setup_styles(doc)
    configure_document(doc)
    bullet_id = add_numbering_definition(doc, ordered=False)
    nav_number_id = add_numbering_definition(doc, ordered=True)
    beat_number_id = add_numbering_definition(doc, ordered=True)
    layer_number_id = add_numbering_definition(doc, ordered=True)
    appendix_number_id = add_numbering_definition(doc, ordered=True)

    # Cover: editorial_cover pattern, with no decorative header rule.
    add_para(doc, "NARRATIVE DESIGN BIBLE", bold_prefix="NARRATIVE DESIGN BIBLE", color=GOLD,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(58)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    set_run_font(title.add_run("《飞花棋》"), size=30, bold=True, color=NAVY)
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.paragraph_format.space_after = Pt(18)
    set_run_font(title2.add_run("剧情设计文档"), size=24, bold=True, color=NAVY)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(8)
    set_run_font(sub.add_run("从“赶考通关”到“借万卷写出自己的心”"), size=14, color=DARK_BLUE)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(74)
    set_run_font(sub2.add_run("世界观 · 主线结构 · 角色弧 · 选择回声 · 结局与文案规范"), size=10.5, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(3)
    set_run_font(meta.add_run("版本 1.0｜内部创作参考"), size=10, bold=True, color=GOLD)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(meta2.add_run("2026-08-22"), size=10, color=MUTED)
    doc.add_page_break()

    add_heading(doc, "阅读约定", 1)
    add_callout(doc, "一句话方向", "玩家不是来背完古人的答案，而是借古人的句子，写出自己的心。")
    add_tagged(doc, "现状", "来自当前 narrative、events、npcs、questions、album、schools、board 与结算文本，可直接视为项目基线。")
    add_tagged(doc, "建议", "为增强连续剧情而新增的设定、角色关系、桥段和字段，允许与当前实现保持一定距离。", GOLD)
    add_tagged(doc, "原则", "后续若与数值或关卡冲突，优先保留主题、人物动机与情绪落点，再调整事件出现条件。", DARK_BLUE)

    add_heading(doc, "内容导航", 2)
    nav = [
        "文档定位——明确这份母版如何指导后续文案",
        "现有剧情基线——盘点已经成立的叙事资产与缺口",
        "核心创意——统一主题、玩家幻想与戏剧问题",
        "世界观规则——解释桃花岛、复合时代与三重科场",
        "主线结构——用十二个关键节拍串起一局",
        "角色设计——让 NPC 成为玩家写作道路的镜子",
        "选择回声——用轻量状态制造“我的故事”",
        "奇遇与支线——建立可持续扩写的内容模版",
        "殿试与桃源终卷——让终局回应整局选择",
        "结局矩阵——基础结局与个性化尾声组合",
        "示例文案——提供可落地的场景与台词样稿",
        "内容生产规范——字段、语体、长度与验收标准",
        "实施路线——按收益与改动成本分批推进",
    ]
    for item in nav:
        add_numbered(doc, item, nav_number_id)

    add_heading(doc, "1. 文档定位", 1)
    add_para(doc, "本文件是一份“剧情母版”，服务于后续奇遇、NPC、阶段过场、题库场景、结算评语和隐藏结局的持续扩写。它不要求游戏立刻实现完整分支，也不把每条建议都视为既定设定。")
    add_heading(doc, "1.1 使用目标", 2)
    goals = [
        "让所有新增文案围绕同一个戏剧问题工作，而不是继续堆叠彼此孤立的典故。",
        "让三种流派成为三种写作人格；玩家选择的是“如何看世界”，不仅是开局属性。",
        "让对手拥有可辨识的欲望与声音，使论战兼具策略信息和人物关系。",
        "用低成本的即时、跨阶段与结局回声，使一次对局形成可复述的个人经历。",
        "保留知识性与古典气质，同时让典故首先服务于现场、冲突和玩家动作。",
    ]
    for g in goals:
        add_bullet(doc, g, bullet_id)
    add_heading(doc, "1.2 不做什么", 2)
    for x in [
        "不强行建立考据严密的单一历史朝代；当前素材跨越先秦至近现代，宜解释为试境中的“文脉同场”。",
        "不把主角写成固定性格与身世；现代记忆保持模糊，只提供方向性压力。",
        "不依赖大规模树状分支；优先采用主线稳定、回声变化、尾声组合的结构。",
        "不让长篇剧情压过棋盘节奏；关键叙事应短、准、可跳过，并能在日志中回看。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "2. 现有剧情基线", 1)
    add_heading(doc, "2.1 已经成立的事实", 2)
    rows = [
        ("主角", "带着模糊的现代记忆抵达桃花岛，被仙人一点后进入科场试境，十年后从童生身份出发。"),
        ("旅程", "棋盘分外圈“起势”、中圈“验收”、内圈“定稿”；功名从秀才、举人、进士逐级收紧。"),
        ("流派", "博闻、奇士、辞宗分别偏向学力、思力、笔力，也代表积累、破局、表达三种写作道路。"),
        ("终局", "常规目标是殿试与金榜；满足全名篇、熟练度和殿试表现后，可进入桃源终圈对战陈之微。"),
        ("结局", "现有结果包括封笔、回合耗尽、殿试、金榜、桃源留问与走出桃源。"),
        ("跨局", "传世名篇、熟练度与“照我传灯”让前局修为留下痕迹，天然适合解释为旧梦与文脉传承。"),
    ]
    add_table(doc, ["要素", "现有文本所表达的基线"], rows, [1800, 7560])
    add_heading(doc, "2.2 最有价值的现有资产", 2)
    for x in [
        "开篇的“现代世界—桃花岛—十年蒙学”构成强钩子，既能容纳玩家投射，也能合理化跨时代典故。",
        "“待到种种妄念破灭，自可殿试见我”已经提前种下隐藏终卷，是主线最重要的承诺。",
        "三圈名称“起势—验收—定稿”同时描述文章成形与人物成长，适合成为整局叙事骨架。",
        "42 个奇遇覆盖苦读、游历、知音、落第、成名、礼佛等人生侧面，具备支线网络的雏形。",
        "NPC 的招牌、破绽与行藏已经能被翻译为性格：守熟路、爱模仿、重稳卷、见势强攻。",
        "选择回声已经为 14 个 choice 奇遇建立了高质量即时反馈，可直接向跨阶段回声扩展。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "2.3 当前缺口", 2)
    add_callout(doc, "核心判断", "现有内容“有文化、有画面、规则清楚”，下一阶段要补的是“有动机、有关系、有回声”。", BLUE)
    for x in [
        "主角为何坚持赴考、为何在金榜之外还要作答，目前主要由系统目标替代。",
        "大多数 NPC 有机制但没有持续关系；玩家打完即忘，无法形成对手、同道或镜像。",
        "奇遇在即时数值上有后果，但少有跨阶段再出现的记忆，选择不易沉淀成人物形象。",
        "殿试与评分能总结构筑，却尚未充分总结玩家的行为与价值倾向。",
        "题库和典故知识层很强，但故事层有时被出处说明压住；玩家像在读注释，而不是在现场。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "3. 核心创意", 1)
    add_heading(doc, "3.1 一句话梗概", 2)
    add_quote(doc, "一个厌倦、失意或迷惘的现代人误入桃花试境，从蒙童一路写到金殿；他借遍古人的声名、与无数文士交锋，最终必须回答：若没有金榜、没有前人的句子，我还会不会写？")
    add_heading(doc, "3.2 核心戏剧问题", 2)
    add_callout(doc, "主问题", "一个人可以借万卷成名，却能否不为万卷所役？")
    for x in [
        "外层问题：我能不能中第，证明自己的能力？",
        "中层问题：我在模仿、求胜与应势之中，是否还认得自己的声音？",
        "内层问题：若功名已得或注定不得，我是否仍愿意落笔？",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "3.3 四个主题支柱", 2)
    pillars = [
        ("功名与自证", "功名不是反派；它给予秩序、压力与见证。危险在于把“被看见”误作“为何写”。"),
        ("传统与自我", "古人之言是路标，不是终点。游戏奖励识典，也应奖励换体、转念和写出自己的句子。"),
        ("才华与代价", "灵感会耗尽，苦读会伤身，成名会催债。真正的成熟包含何时燃笔、何时收卷。"),
        ("独行与知音", "文章看似独作，实由师友、对手、读者和无名之人共同照见；“与人相遇”应成为成长来源。"),
    ]
    add_table(doc, ["主题", "剧情表达"], pillars, [2200, 7160])
    add_heading(doc, "3.4 玩家幻想", 2)
    for x in [
        "我能在古典文脉中自由行走，认出典故，也能把典故变成自己的选择。",
        "我逐渐形成一套独特写法；对手会记住、模仿、质疑甚至敬畏它。",
        "我的失败不是清零，而是成为下一局的一页残卷、一点灯火或一条未答完的问题。",
        "最高成就不是被系统宣布“最强”，而是走出棋盘后仍知道下一步为何而走。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "4. 世界观规则", 1)
    add_heading(doc, "4.1 桃花试境", 2)
    add_tagged(doc, "建议", "把桃花岛明确为“由古今文章、人物传说与考场妄念叠成的试境”。岛并非历史朝代，也不是死后世界，而是一处只在求问者心中成立的文心之境。", GOLD)
    add_para(doc, "陈之微能点碎景物、重组成蒙学馆；棋盘会把十年寒窗压缩为一局；前世修为可以由“照我传灯”传下。这些现象都说明世界的基本法则不是物理，而是“文字、记忆与执念可以成形”。")
    add_heading(doc, "4.2 复合时代规则", 2)
    for x in [
        "先秦典籍、唐诗、宋词、明清掌故乃至近现代作品可以同场出现，因为试境按文脉而非年代组织材料。",
        "历史人物原则上不以本尊长期登场；更适合通过梦、碑、残卷、传闻、别号与题目被召回。",
        "具名 NPC 是试境中的考生与考官，不必对应真实人物；他们可继承某种传统，但不能冒充历史本人。",
        "若引用近现代语句，必须确认语体与版权边界；主叙事优先使用公版典籍或原创表达。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "4.3 三重科场的象征", 2)
    stage_rows = [
        ("起势｜外圈", "学会使用已有规则", "“我能写什么？”", "蒙学、初战、第一次选择与第一次被记住"),
        ("验收｜中圈", "发现熟路也会反噬", "“这真是我的写法吗？”", "模仿者、公开评价、选择回返与第一次主动换路"),
        ("定稿｜内圈", "承担自己写法的代价", "“我愿为哪一句负责？”", "名望、文债、权威审视与取舍收束"),
        ("无字｜终圈", "放下题目与排名", "“无人命题时，我还写吗？”", "桃径、无字卷、陈之微与离开棋盘"),
    ]
    add_table(doc, ["层级", "成长任务", "阶段之问", "典型内容"], stage_rows, [1600, 2200, 2200, 3360])
    add_heading(doc, "4.4 反复出现的意象", 2)
    for x in [
        "笔：借来的才华、自己的手艺、燃尽的代价。五色笔与普通毛笔应形成对照。",
        "灯：学习、传承、孤独与有限的时间。“照我传灯”可成为跨局叙事主意象。",
        "水：路途、时间与隔岸的知音。船、渡口、江雨天然承担转场。",
        "榜：被看见的公共尺度。榜纸可在落第、题名与终局反复出现。",
        "桃花：试境的缝隙。常规阶段只偶尔逆风出现，终圈才形成明确道路。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "5. 主线结构", 1)
    add_callout(doc, "结构策略", "保持棋盘主循环不变，用“开场一钩、每圈三拍、殿试两拍、终圈两拍”组成十二个关键节拍。普通格子继续承担自由游历。")
    beats = [
        ("旧梦有缺：", "开场闪回三个互不相容的现代人生版本，强调主角连自己为何来桃花岛都记不真切。"),
        ("仙人点额：", "陈之微只提出条件，不解释答案：“妄念尽时，来交最后一卷。”玩家选择流派，即选择第一种自我解释。"),
        ("第一次被看见：", "与周小满等童生交锋。对手的具体反应让玩家意识到：写法会在别人身上留下痕迹。"),
        ("选择留下墨痕：", "首个 choice 奇遇除即时回声外，记录一个隐藏倾向；日志以一句短评保存。"),
        ("秀才门前：", "第一次晋阶不是单纯升级。先生或旧同窗问：“你赢的是他，还是赢了昨日的自己？”"),
        ("镜中之敌：", "会试圈优先安排苏明哲等模仿型 NPC；他复制玩家近两场路数，逼玩家意识到“套路可被他人占用”。"),
        ("纸贵之后：", "中段安排成名或公开评价事件。赞誉带来收益，也触发文债、期待或自我重复的压力。"),
        ("旧选回返：", "一位曾受帮助的乡人、旧友或未收下的经书再次出现，用一句话改变当下场景，不一定改变数值。"),
        ("定稿之门：", "进入内圈前，玩家看到本局三条“行卷评语”：常用文体、最常见倾向、最难忘的选择。"),
        ("金殿三问：", "殿试仍可只打一场，但由王侍郎、李学士、赵大儒依次追问变通、形式与责任；题面由本局行为生成。"),
        ("金榜之外：", "金榜已定，桃花逆风而开。进入终圈不再增加常规名次，只回答开篇承诺。"),
        ("无字终卷：", "陈之微不给题目。胜负判断表面仍来自构筑，叙事上则是玩家能否在没有规定答案时完成一卷。"),
    ]
    for title, body in beats:
        add_numbered(doc, f"{title}{body}", beat_number_id, title=title)

    add_heading(doc, "5.1 每圈的情绪节奏", 2)
    for x in [
        "起势：好奇 → 初胜 → 代价初现。文案明亮、具体，允许童生的笨拙与幽默。",
        "验收：熟练 → 被模仿 → 自我怀疑 → 主动换路。场景更公开，评价更多来自他人。",
        "定稿：名望 → 催逼 → 取舍 → 承担。句子更短，留白更多，考官少解释。",
        "终圈：安静 → 回望 → 无题 → 离开或留问。减少典故，让原创意象接管叙事。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "6. 角色设计", 1)
    add_callout(doc, "角色原则", "NPC 不是一张会说话的属性卡，而是对某种写作信念的坚持者。玩家破掉他的机制，也是在动摇他的信念。", BLUE)
    add_heading(doc, "6.1 核心角色", 2)

    add_character(doc, "玩家｜失去来历的求问者", "留白主角", "在金榜上证明自己没有白来。", "只要写得足够好，迷惘就会消失。", "把外界评价当成内心答案；现代记忆越模糊，越依赖功名确认存在。", "从借身份、借典故、借排名确认自己，走向承认“我为何落笔”只能由自己回答。", "始终以第二人称表现，不替玩家宣告单一性格；由选择倾向和常用文体拼出轮廓。", [
        "开局内心句：你记不清自己为何出海，却记得那时有一种日子，再过一天也不会不同。",
        "定稿内心句：榜在前面，路也在前面；第一次，你发现两者未必是同一个方向。",
    ], bullet_id)

    add_character(doc, "陈之微｜桃花仙人", "门槛守望者 / 最终对手", "让求问者自行走出试境，而不是被他送出去。", "答案一旦由他人告知，就仍是借来的答案。", "他的冷静容易被误读为冷漠；过早抽走解释，会让人只剩执念。", "从只设题不干预，到在终卷承认玩家已经能为自己的路负责。", "短句、停顿、少赞美；不解释象征，不用高高在上的玄学堆砌。", [
        "登场：你终于来了。先别问这是不是最后一题。",
        "终卷：金榜是别人替你写的名字。此卷无题，只看你还肯写什么。",
        "败后：你已能胜人，只是还未尽胜旧日之己。此问不必今日作答。",
    ], bullet_id)

    add_character(doc, "周小满｜蒙学童子", "第一面镜子 / 可成长的旧识", "证明年纪小也能写出让人记住的诗。", "天分就是能第一个写出来。", "怕暴露声律不稳，常用抢先与嘴硬掩饰。", "从把玩家当“大人”仰望，到在后续来信中以自己的笨拙方式坚持写作；他证明成长不只发生在主角身上。", "直白、快、藏不住情绪；禁止老成官话。", [
        "登场：周小满把卷子抱在胸前，先看你的笔，又偷偷看了眼自己的。",
        "被破招：你怎么换了写法？……再来一篇，我这回听得出来。",
        "后续来信：我还是常把平仄写反。先生说，写反了也先别把纸烧掉。",
    ], bullet_id)

    add_character(doc, "苏明哲｜公车举人", "镜像对手", "证明自己能比任何人更快掌握“正确写法”。", "好文章都有可复用的法门；看得够准，就不必冒险。", "模仿越成功，越恐惧别人问他自己想说什么。", "前期以复制玩家为压迫，中后期在玩家主动换路时短暂失措；若再次相遇，他会第一次不用玩家的常用体。", "礼貌、圆熟、常复述对方的话；情绪从不直接承认。", [
        "登场：苏明哲翻过你近两场的行卷，笑道：“路既走得通，多走一遍又何妨？”",
        "被破招：这一笔不在你旧卷里。……也好，今日总算不是照着影子写。",
        "惜败：他收卷很慢，末了问：“若不用你最擅长的那一体，你还认得自己吗？”",
    ], bullet_id)

    add_character(doc, "唐季卿｜孝廉举人", "传统的辩护者", "守住经义与文脉，不让科场变成只争奇巧的戏台。", "先有共同的经典，才有可以互相理解的文章。", "把“熟”误作“真”，难以承认亲历也能修正经注。", "在老农问字、古寺残碑等回声影响下，看见学问若不能回到人间，就只是卷中自转。", "严谨、完整、爱引用；被说服时先沉默，再承认一个很小的例外。", [
        "登场：唐季卿把经注压在卷角：“新意可以有，先说你从哪一页走来。”",
        "回声：你曾替田埂上的人讲明官文。唐季卿看了许久，只说：“经义原也该有人听懂。”",
    ], bullet_id)

    add_character(doc, "宇文渊｜甲科进士", "立意审判者", "写出足以压过辞采的中心判断。", "文章首先要知道自己为何存在。", "轻视形式与人情，容易把复杂经验压成一句结论。", "他迫使玩家说出立意，玩家则迫使他承认有些真意必须通过细节和情感才能抵达。", "冷、准、提问多于陈述；从不为漂亮句子单独喝彩。", [
        "登场：宇文渊没有看你的起句，只以指节敲了敲末段：“你究竟要人信什么？”",
        "被破招：他第一次回头看了你的景物描写：“原来这层意思，直说反而浅了。”",
    ], bullet_id)

    add_character(doc, "王侍郎｜礼部侍郎", "殿试主审 / 制度化权威", "从会写的人中挑出能在变化中自持的人。", "真正成熟的路数，必须经得起被看穿。", "过于相信观察和适应，容易把人的改变也当成可计算的策略。", "他会记住玩家旧路；若玩家的变化来自真实选择而非临场投机，他最终给出“此卷有主”的评语。", "克制、官式但不空泛；用批注、停笔和反问施压。", [
        "首问：前两场皆以一锋破题。今日还要走旧路么？",
        "换体后：知道换笔，不难。换了笔，意思还在，才难。",
        "评卷：套路可变，所守未变。此卷有主。",
    ], bullet_id)

    add_heading(doc, "6.2 角色投放原则", 2)
    for x in [
        "每局固定 2—3 名“关系 NPC”跨阶段出现，其余仍可作为随机对手，避免重复内容过重。",
        "关系 NPC 的第二次出现必须读取一项可感知行为：常用文体、曾触发的奇遇、是否主动换体或最近胜负。",
        "同一 NPC 的强度升级可解释为“对方也在赶考”，而不是复制一个更高数值版本。",
        "胜负台词都要维护人物尊严；童生可以沮丧，进士可以冷淡，但不要让失败者变成笑料。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "7. 选择回声系统", 1)
    add_heading(doc, "7.1 四条隐藏倾向", 2)
    add_para(doc, "建议不在 HUD 上显示道德条，也不判定“正确答案”。每次 choice 只为一到两个倾向轻量记分；它们用于选取日志、NPC 评语和尾声，不直接制造强数值最优解。")
    inclinations = [
        ("逐名 ↔ 求真", "追求被看见与承认", "追问作品是否诚实、是否属于自己"),
        ("守法 ↔ 出新", "敬畏规范、典籍与稳定", "换体、转念、破格与自造"),
        ("与人 ↔ 独行", "授业、知音、回应与公共责任", "保留、远行、自守与内在完成"),
        ("惜身 ↔ 燃笔", "知道停笔、留力、来日再写", "以灵感、健康或安稳换取当下成篇"),
    ]
    add_table(doc, ["轴线", "左端含义", "右端含义"], inclinations, [1900, 3730, 3730])
    add_heading(doc, "7.2 四层回声", 2)
    echoes = [
        ("即时", "选择后 1 个界面内", "resultText", "兑现动作与现场余韵；当前 14 个 choice 已具备"),
        ("短程", "同圈 6—18 格内", "日志 / 路人格", "某物件、路人或一句旧话再次出现"),
        ("跨阶段", "下一圈门槛或关系 NPC", "过场 / 台词", "把选择解释为玩家逐渐形成的习惯"),
        ("终局", "殿试评语与尾声", "epilogue token", "不复述数值，而是回答“你成了怎样的写作者”"),
    ]
    add_table(doc, ["层级", "时机", "载体", "作用"], echoes, [1200, 1900, 1900, 4360])
    add_heading(doc, "7.3 四条示例链", 2)
    chains = [
        ("E006 江郎才尽｜还笔", "即时：五色光退去，第一笔很慢。", "短程：普通毛笔在战后断裂，却留下最稳的一行。", "殿试：王侍郎批“辞不借色，意自有光”。", "尾声：你未带走五色笔，只带走一支用旧的笔。"),
        ("E018 老农问字｜停步讲解", "即时：乡邻能念出要紧处。", "跨阶段：唐季卿承认经义也该有人听懂。", "殿试：赵大儒追问“此卷能否出金殿、到田埂？”", "尾声：榜外有人请你再讲一遍，你停下了脚。"),
        ("E025 落第榜下｜写愁", "即时：榜名未变，纸上留下不肯沉默的一笔。", "短程：陌生客在舟中吟出其中半句。", "殿试：李学士问“你要写给失意者，还是写给评卷者？”", "尾声：金榜之外，那首无名诗先被人记住。"),
        ("E042 留人古寺｜辞经", "即时：松风重复“因果未尽”。", "跨阶段：定稿门前，包袱里出现一片旧经书签。", "终圈：桃径旁传来同一声寺钟。", "尾声：你没有带走经书，却终于读懂拒绝本身也是一页。"),
    ]
    for title, immediate, later, final, ending in chains:
        add_heading(doc, title, 3)
        for item in (immediate, later, final, ending):
            add_bullet(doc, item, bullet_id)

    add_heading(doc, "7.4 回声选择规则", 2)
    for x in [
        "同一条选择最多安排一次强回声与一次弱回声，避免玩家感觉每件事都被系统监视。",
        "强回声优先给具备人物、物件和动作的事件；纯资源选择可以只保留即时回声。",
        "回声不反转玩家原意。选择休息，不应在后文被嘲讽为怯懦；选择燃笔，也不自动等于高尚。",
        "多个回声竞争时，优先级为：关系 NPC 关联 > 本圈最近选择 > 稀有事件 > 随机通用句。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "8. 奇遇与支线设计", 1)
    add_heading(doc, "8.1 六类内容池", 2)
    event_types = [
        ("行路", "雨、舟、长亭、客舍、驿馆", "调节节奏，承接转场与乡愁"),
        ("治学", "抄书、残碑、推敲、灯油", "呈现功夫与代价，连接学力/思力/笔力"),
        ("交游", "知音、诗会、卖字、问字", "建立公共世界与关系回声"),
        ("名望", "纸贵、题名、赐宴、被誉", "把成长转化为期待、催债与诱惑"),
        ("失意", "落第、病中、才尽、败后", "让失败成为创作来源，而非纯惩罚"),
        ("超越", "桃花、古寺、梦笔、无字", "少量投放，推进陈之微与终圈伏笔"),
    ]
    add_table(doc, ["类型", "现有意象", "剧情职责"], event_types, [1500, 3000, 4860])
    add_heading(doc, "8.2 单张奇遇的三层结构", 2)
    for title, body in [
        ("现场层：", "先写玩家此刻看见的物件与正在发生的动作。"),
        ("压力层：", "说明为什么必须现在选择；时间、体力、名声、他人需要或自我诱惑至少占一项。"),
        ("文脉层：", "典故作为照面、反差或记忆钩子；出处移到次级区域，不替玩家解释人生道理。"),
    ]:
        add_numbered(doc, f"{title}{body}", layer_number_id, title=title)
    add_heading(doc, "8.3 支线组合方式", 2)
    for x in [
        "物件链：五色笔 → 普通笔 → 断笔 → 终卷旁的桃枝；适合跨阶段而不依赖同一 NPC。",
        "人物链：周小满、苏明哲、唐季卿各自出现 2—3 次，第二次读取前一次胜负或选择。",
        "命题链：同一主题先以知识题出现，再以 choice 出现，最后由考官追问价值判断。",
        "名篇链：传世名篇解锁时，不只弹出奖励；增加“一页如何被人传抄”的微型尾声。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "8.4 示例：把知识题变成剧情资产", 2)
    add_para(doc, "题库中“同一场景可以通向不同写法”的 choice 题，是剧情化最成熟的一组。建议为每题记录玩家选择的 attr，并在后续生成一句“行卷气口”：例如重阳登高选择抱负，定稿评语可写“你总把高处当作向前的证据”；选择思念，则写“你登得越高，越记得缺席的人”。")

    add_heading(doc, "9. 殿试与桃源终卷", 1)
    add_heading(doc, "9.1 殿试：三问一战", 2)
    add_para(doc, "当前同心棋盘以殿试单场取胜判定金榜，而 NPC 池已有三位主考官。剧情上可把三人改造成同一场考试的三个审视角度：三问改变题面、风潮或评语，最后只进行一次正式论战。")
    exam_rows = [
        ("王侍郎｜变", "你是否被自己的常用路数困住？", "读取近两场文体、是否换体、是否追加骰", "衡文察变"),
        ("李学士｜情", "形式与声律之外，这篇文章究竟在意谁？", "读取常用文风、知音/送别/落第类选择", "殿试声律"),
        ("赵大儒｜用", "此卷若出金殿，能否被榜外之人听懂？", "读取问字、授业、经义与守法/出新倾向", "经义稳卷"),
    ]
    add_table(doc, ["考官", "所问", "读取", "现有机制锚点"], exam_rows, [1800, 2800, 2960, 1800])
    add_heading(doc, "9.2 行为化评语", 2)
    for x in [
        "连续同体胜利：前两卷皆以一锋破题。锋可用，亦可反过来役使执笔之人。",
        "主动换体并获胜：知道换笔，不难；换了笔，意思还在，才难。",
        "高频追加灵感骰：卷有急气。敢燃笔是胆量，知道哪一句值得燃，才是分寸。",
        "多次选择助人：你的文章常为别人停步。金殿问的是功名，榜外问的是用处。",
        "多次选择休息：能在该停时收卷，并非怯弱。只是今日这一卷，不可留到明日。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "9.3 隐藏终圈：从有题到无题", 2)
    add_para(doc, "隐藏终圈的十二格“桃门、花影、忘机、涤心、照水、听泉、无字、回梦、问真、见性、归一、终卷”已经自带仪式路径。建议不再塞入普通奇遇，而用极短回望句，让玩家逐步卸下系统身份。")
    secret_lines = [
        "桃门：门上没有匾，也没有榜。",
        "花影：花影掠过衣袖，像有人翻过你第一场的旧卷。",
        "忘机：骰声停了。你忽然听见自己的呼吸。",
        "照水：水里先映出童生冠，继而是现代衣装，最后都被涟漪打散。",
        "无字：案上只有一张纸。没有题目，也没有倒计时。",
        "回梦：你想起出海前那种“一日与一日没有不同”的感觉。",
        "问真：若不为中第，你还剩哪一句要写？",
        "归一：诗、词、联、笔、学、思的名字暂时退去，只剩一次落笔。",
    ]
    for line in secret_lines:
        add_bullet(doc, line, bullet_id, bold_prefix=line.split("：")[0] + "：")
    add_heading(doc, "9.4 终卷胜负的叙事解释", 2)
    add_tagged(doc, "胜", "不是打败仙人，而是在没有命题与名次时仍能完成一卷；陈之微因此不再替玩家标路。", DARK_BLUE)
    add_tagged(doc, "负", "不是被剥夺金榜，而是仍需要外在尺度才能定稿；桃源保留问题，允许下一局继续作答。", GOLD)

    add_heading(doc, "10. 结局矩阵", 1)
    add_para(doc, "采用“基础结局 + 流派落款 + 倾向尾声 + 关系 NPC 片段”的组合方式。基础结局控制事实，后置片段提供个性；无需为所有组合手写完整分支。")
    ending_rows = [
        ("封笔｜江郎才尽·悔", "灵感耗尽", "你停在途中，但未否定曾写下的东西", "下一局由残卷、旧笔或传灯续接"),
        ("岁月不居", "回合上限仍未抵达", "不是失败于才力，而是无法完成取舍", "尾声强调未定稿，不嘲讽拖延"),
        ("殿试已毕", "抵达但未夺魁", "权威给出评价，玩家仍保有自己的行卷", "按最高评分维度生成去向"),
        ("金榜题名", "殿试取胜", "功名愿望兑现，但开篇之问尚未必回答", "若未入终圈，保留桃花伏笔"),
        ("桃源留问", "终卷失利", "金榜不被收回，终问留待来局", "最适合连接跨局传承"),
        ("走出桃源", "终卷取胜", "不再由棋盘标路，第一次主动选择去向", "以倾向决定离岛后的第一件事"),
    ]
    add_table(doc, ["基础结局", "条件", "情绪落点", "后续接口"], ending_rows, [2000, 1750, 3000, 2610])
    add_heading(doc, "10.1 流派落款", 2)
    for x in [
        "博闻：你带走的不是更多书，而是知道何时合上书，听眼前人说完一句。",
        "奇士：路离开棋盘后没有格子。你反而第一次确定，这一步不是走错。",
        "辞宗：你没有再等一句更好的起笔。风吹动纸角时，第一行已经落下。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "10.2 倾向尾声样例", 2)
    for x in [
        "逐名偏高：你仍走向人多处。不同的是，这次榜上有没有名字，不再决定你是否开口。",
        "求真偏高：你把最得意的一篇留在岛上，只带走那篇尚未写完的。",
        "与人偏高：下山第一日，你在田埂边停住，为人念完一张官文。",
        "独行偏高：你没有回京赴宴，沿江走了很久；无人催稿，句子来得比从前慢。",
        "惜身偏高：夜深时你准时熄灯。那页书夹着枯叶，明日仍可翻开。",
        "燃笔偏高：你仍会为一句话忘记天亮，只是不再把耗尽自己当成才华的证明。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "11. 示例文案", 1)
    add_heading(doc, "11.1 序章建议稿", 2)
    add_quote(doc, "你对来处只剩三种互相矛盾的记忆：曾拥有得太多，曾一无所有，也曾与世上多数人一样，把一日过成另一日。\n\n它们都在海上断了。雾散时，桃花岛浮在水面，岸边有人问你：“来求什么？”\n\n你说不清。那人便以指点在你额上。山海碎成万点墨色，再聚拢时，你已坐在蒙学馆最末一席，手里握着一支寻常毛笔。\n\n“待种种妄念破灭，自可殿试见我。”\n\n十年后，科举将启。你仍不知那段旧梦是真是假，只知道这一次，卷纸已经铺开。", "建议稿｜保留现有设定并强化“寻求答案”的动机")
    add_heading(doc, "11.2 阶段过场", 2)
    add_heading(doc, "秀才门前｜第一次晋阶", 3)
    add_quote(doc, "县学门前贴出新榜。周小满踮着脚替你找名字，找着后比你先笑起来。\n\n“以后再比，你可不许让着我。”\n\n你把名字看了两遍。纸上的墨很新，心里那句“只要中第便会明白”却没有因此更清楚。")
    add_heading(doc, "会文门｜进入验收", 3)
    add_quote(doc, "外圈一路积下的卷子被重新装订。翻到最熟的一篇时，你几乎能背出下一处转折。\n\n门内有人也在念同样的句子，连停顿都与你相似。\n\n苏明哲合上你的旧卷，向你一揖：“路既走得通，多走一遍又何妨？”")
    add_heading(doc, "登第门｜进入定稿", 3)
    add_quote(doc, "门吏不看你的名帖，只递回三张行卷：一张写你最常用的文体，一张记你最重的一次取舍，最后一张空白。\n\n“前两张是你走来的路。”他说，“第三张，留给你今日愿意负责的那一句。”")
    add_heading(doc, "11.3 NPC 交锋样稿：苏明哲", 2)
    add_quote(doc, "苏明哲把你近两场行卷平码在案上。你的起句、转折与收束，被他用另一种笔迹重新写了一遍。\n\n“文章既有法，何必每回都冒险？”\n\n本场行藏：他将优先模仿你近日常用路数。若你主动换体，他的招牌会失去依托。", "登场与机制提示合并示例")
    add_heading(doc, "11.4 旧选回返样稿：老农问字", 2)
    add_quote(doc, "入会文门前，一名差役追上来，递给你半张折得发软的纸。\n\n纸上只有几行歪斜的大字，是那日田埂上的乡人合写的：官文已经听懂，田租也问明白了。末尾空着一格，像是等你回一句。\n\n唐季卿在旁看了许久，低声道：“经义原也该有人听懂。”")
    add_heading(doc, "11.5 殿试问答样稿", 2)
    add_quote(doc, "王侍郎：前两卷皆以诗体争先。今日还走旧路？\n\n李学士：你句中有风月。风月之外，究竟在意谁？\n\n赵大儒：此卷若传出金殿，榜外之人可听懂一句么？\n\n主考官搁下朱笔：不必答我们。落笔便是你的答复。")
    add_heading(doc, "11.6 走出桃源｜扩写样稿", 2)
    add_quote(doc, "陈之微看完终卷，没有判等，也没有落印。\n\n“金榜是别人替你写的名字。此卷，才是你自己留下的。”\n\n漫天花影碎开，童生铺、三重科场与金殿都缩成水面的一点灯火。你接过那枝桃花，回身时，脚下再没有方格。\n\n远处有一条路，也可能没有。你先走了一步。")

    add_heading(doc, "12. 内容生产规范", 1)
    add_heading(doc, "12.1 语体", 2)
    for x in [
        "玩家视角统一第二人称，但避免连续三句都以“你”开头；优先让物件、声音和动作带出现场。",
        "叙述用清晰现代白话承载古典意象；不追求通篇仿古，不让生僻词遮住选择。",
        "童生直白，秀才讲体面，举人会试探，进士克制精确，考官少解释；去掉名字后仍应能分辨至少两类声音。",
        "禁止手机、社交媒体、后台配置等穿透场景的现代词；现代记忆只以感受、物件轮廓或不可靠闪回出现。",
        "引号统一中文弯引号；范围连接使用全角破折号或“至”，避免混用半角符号。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "12.2 长度预算", 2)
    length_rows = [
        ("普通奇遇正文", "60—100 字", "一物、一动、一压力"),
        ("稀有奇遇正文", "80—130 字", "可多一层人物或典故反差"),
        ("传说奇遇正文", "100—160 字", "超出则拆为主文案与典故层"),
        ("选择项", "12—28 字", "两项长度和文学密度尽量接近"),
        ("即时回声", "35—70 字", "只写现场后果，不解释数值"),
        ("NPC 登场", "35—80 字", "动作 + 一句人声 + 必要机制"),
        ("阶段过场", "90—180 字", "一件事发生 + 一个问题推进"),
        ("结局尾声", "120—260 字", "事实收束 + 意象回扣 + 下一步"),
    ]
    add_table(doc, ["文本类型", "建议长度", "必须完成的任务"], length_rows, [2400, 1800, 5160])
    add_heading(doc, "12.3 奇遇字段建议", 2)
    code = '''{
  "id": "E0XX",
  "name": "奇遇名",
  "sceneText": "现场层：物件、动作与压力",
  "allusion": {"source": "出处", "note": "次级说明"},
  "choices": [{
    "text": "玩家动作",
    "resultText": "即时回声",
    "echoTags": ["truth", "fellowship"],
    "echoKey": "farmer_letter",
    "effect": {}
  }]
}'''
    p = doc.add_paragraph(style="Story Code")
    shade_paragraph(p, LIGHT_GRAY, DARK_BLUE)
    set_run_font(p.add_run(code), size=8.5, color=INK, ascii_font="Consolas", cjk_font="Microsoft YaHei UI")
    add_heading(doc, "12.4 NPC 字段建议", 2)
    code = '''{
  "intro": "登场动作与人声",
  "winLine": "NPC 获胜后的反应",
  "loseLine": "NPC 失利后的反应",
  "weaknessLine": "破绽被击中时的反应",
  "memoryLines": {
    "farmer_letter": "读取旧选的定向台词",
    "switched_style": "读取玩家行为的定向台词"
  },
  "voice": {"keywords": ["克制", "反问"], "avoid": ["夸张挑衅"]}
}'''
    p = doc.add_paragraph(style="Story Code")
    shade_paragraph(p, LIGHT_GRAY, DARK_BLUE)
    set_run_font(p.add_run(code), size=8.5, color=INK, ascii_font="Consolas", cjk_font="Microsoft YaHei UI")
    add_heading(doc, "12.5 史料与典故政策", 2)
    for x in [
        "先核对作者、篇名、时代、原句与典故来源；不确定时宁可删去具体归属，也不要写成确定事实。",
        "正文中的古典原句只承担一个功能：照面、对比或记忆钩子；完整解释进入折叠层。",
        "争议性掌故应使用“相传”“旧说”等限定语，不把文学传说当作无争议史实。",
        "同一内容包内避免重复高频名句；每批新增文案应统计作者、朝代、主题与意象分布。",
        "知识题解析先给答案与记忆钩子，再说明干扰项；剧情选择题不标标准答案。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "12.6 单条文案验收", 2)
    checks = [
        "不看典故出处，玩家也知道此刻发生了什么。",
        "至少有一个可见物件和一个正在发生的动作。",
        "choice 的两项代表不同取向，而非明显的高奖励与低奖励。",
        "结果回声兑现玩家动作，不替玩家总结人格。",
        "NPC 台词符合档位、性格与当前胜负，不只是机制说明换皮。",
        "文本在移动端长度预算内，关键信息不埋在最后一句。",
        "若读取旧选择，回声与原意一致，且玩家能认出关联。",
    ]
    for c in checks:
        add_bullet(doc, f"□ {c}", bullet_id)

    add_heading(doc, "13. 实施路线", 1)
    roadmap = [
        ("P0｜建立主线可见性", "低", "序章、三阶段过场、殿试行为评语、五类基础结局", "玩家能说清为何出发、为何继续、终局回答了什么"),
        ("P0｜关系 NPC 首批", "低—中", "周小满、苏明哲、唐季卿、王侍郎各 3—5 条动态台词", "连续遇见时可凭声音与记忆区分角色"),
        ("P1｜跨阶段回声", "中", "为 14 个 choice 选出 6 个强回声，建立 echoTags/echoKey", "至少半数通关局出现 2 次可识别旧选回返"),
        ("P1｜殿试三问一战", "中", "三考官各读取 1—2 项行为并生成题前评语", "评语不复述数值，能准确描述玩家本局习惯"),
        ("P2｜支线网络", "中—高", "新增物件链、人物链、命题链与名篇微尾声", "随机内容仍有变化，但一局可形成 2—3 条完整小弧"),
        ("P2｜个性化尾声", "中", "基础结局叠加流派、倾向与关系 NPC 片段", "同一名次的两局可读出明显不同的人物画像"),
    ]
    add_table(doc, ["批次", "改动", "交付", "验收重点"], roadmap, [1800, 1200, 3260, 3100])
    add_heading(doc, "13.1 推荐首批内容量", 2)
    for x in [
        "4 段主线过场：序章、会文门、登第门、殿试前。",
        "4 名关系 NPC：每人登场、胜、负、破绽、旧选读取各 1 条，合计约 20 条。",
        "6 条强回声：从江郎才尽、老农问字、落第榜下、留人古寺、知音、焚膏继晷中选取。",
        "15 条殿试评语：按文体、换体、追加骰、助人、休息、燃笔等行为组合。",
        "5 个基础结局 + 3 个流派落款 + 8 个倾向尾声，采用片段拼接而非 120 个整段分支。",
    ]:
        add_bullet(doc, x, bullet_id)
    add_heading(doc, "13.2 叙事效果指标", 2)
    for x in [
        "首次完成一局后，玩家能用一句话复述自己的关键选择，而不只记得最终分数。",
        "测试者能说出至少两名 NPC 的性格或信念，而不只说“诗体敌人”“模仿机制”。",
        "殿试评语至少有一条让玩家明确意识到系统记得自己前面的行为。",
        "隐藏终圈被理解为主题终局，而不只是更难的一场 Boss 战。",
        "典故知识点仍清晰，但删去出处层后，主场景也能独立成立。",
    ]:
        add_bullet(doc, x, bullet_id)

    add_heading(doc, "附录 A｜现有素材映射", 1)
    source_rows = [
        ("narrative.json", "序章、风潮、阶段晋阶、会试再入、隐藏终圈邀请/胜/负", "主线事实与终局原句"),
        ("events.json", "42 个奇遇；14 个 choice 含 28 条 resultText", "支线池、选择回声与主题倾向"),
        ("npcs.json", "童生至主考官、桃花仙人；招牌/破绽/行藏", "人物信念与动态台词锚点"),
        ("questions.json", "49 道知识题、18 道风格选择题", "命题场景、价值选择与写作人格"),
        ("schools.json", "博闻、奇士、辞宗的 motto/flavor/desc", "三种主角起始人格"),
        ("album.json", "12 篇传世名篇、跨局成长与双路线", "跨局记忆、名望与个人传承"),
        ("board.json", "起势/验收/定稿三圈与桃源终圈十二格", "空间化叙事节拍"),
        ("grades.json / game.js", "评分维度、功名等第与基础结局原因", "结算事实与尾声组合条件"),
        ("既有评审与回声终稿", "文案问题诊断、14 个选择事件的即时回声", "写作规范与已完成成果"),
    ]
    add_table(doc, ["素材", "主要内容", "本文件用途"], source_rows, [2200, 4200, 2960])
    add_heading(doc, "附录 B｜后续创作的三个判断题", 1)
    for q in [
        "这段文字是在讲古人的故事，还是在推进玩家此刻的故事？如果两者冲突，优先后者。",
        "这个 NPC 坚持什么，玩家的写法为什么会伤到或改变他？如果答不出，机制还没有变成人物。",
        "这项选择会留下什么可被再次看见的东西？如果只有数值，至少补一个即时回声。",
    ]:
        add_numbered(doc, q, appendix_number_id)
    add_callout(doc, "最终准绳", "古人的句子是桥，玩家自己的那一步才是彼岸。", GOLD)

    doc.core_properties.title = "《飞花棋》剧情设计文档"
    doc.core_properties.subject = "世界观、主线结构、角色弧、选择回声、结局与文案规范"
    doc.core_properties.author = "飞花棋项目组"
    doc.core_properties.keywords = "飞花棋, 剧情设计, 文案, Narrative Design"
    doc.core_properties.comments = "基于 2026-08-22 工作区现有文本整理。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
