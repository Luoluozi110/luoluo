from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "《文心棋》玩家介绍.docx"
COVER_ART = ROOT / "feihuaqi-playable" / "assets" / "art" / "peach-academy-island-v1.png"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
INK = "273A32"
CRIMSON = "8A3B2E"
JADE = "245C4F"
GOLD = "9A6C25"
MUTED = "6B706C"
PAPER = "FBF7ED"
PALE_JADE = "EAF1EC"
PALE_GOLD = "F7EEDB"
PALE_RED = "F5E8E3"
LINE = "D8D0C0"


def set_run_font(run, size=None, bold=None, color=None, italic=None, east_asia=FONT_CN):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), FONT_LATIN)
    rfonts.set(qn("w:hAnsi"), FONT_LATIN)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    trpr.append(cant_split)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color=LINE, size=6):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(size))
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border_left(paragraph, color=CRIMSON, size=18, space=8):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def remove_paragraph_borders(paragraph_or_style):
    element = paragraph_or_style._p if hasattr(paragraph_or_style, "_p") else paragraph_or_style._element
    ppr = element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is not None:
        ppr.remove(pbdr)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, page_text, fld_end])
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=MUTED)


def set_picture_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = FONT_LATIN
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    remove_paragraph_borders(title)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT_LATIN
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(JADE)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 16, CRIMSON, 18, 10),
        ("Heading 2", 13, JADE, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_custom_numbering(doc, num_format, text, color=CRIMSON):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=-1) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), FONT_CN)
    rpr.append(rfonts)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    rpr.append(col)
    lvl.append(rpr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_body(doc, text, *, bold_lead=None, align=None, after=6, keep=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 11, bold=True, color=CRIMSON)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, color=INK)
    return p


def add_labeled(doc, label, text, *, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    r1 = p.add_run(label)
    set_run_font(r1, 11, bold=True, color=CRIMSON)
    r2 = p.add_run(text)
    set_run_font(r2, 11, color=INK)
    return p


def add_bullet(doc, num_id, text, bold_lead=None):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, 11, bold=True, color=JADE)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, 11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, 11, color=INK)
    return p


def add_callout(doc, title, text, fill=PALE_GOLD, border=GOLD):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, fill)
    set_paragraph_border_left(p, border)
    r1 = p.add_run(title)
    set_run_font(r1, 11, bold=True, color=border)
    r2 = p.add_run(text)
    set_run_font(r2, 11, color=INK)
    return p


def add_numbered_step(doc, num_id, title, text, after=4):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_together = True
    r1 = p.add_run(title)
    set_run_font(r1, 11, bold=True, color=CRIMSON)
    r2 = p.add_run(text)
    set_run_font(r2, 11, color=INK)
    return p


def add_section_title(doc, title, eyebrow=None):
    if eyebrow:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(eyebrow)
        set_run_font(r, 9.5, bold=True, color=GOLD)
    p = doc.add_paragraph(title, style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    return p


def add_footer_and_header(section):
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run("文心棋  |  玩家介绍")
    set_run_font(run, 9, bold=True, color=MUTED)
    fp = section.footer.paragraphs[0]
    add_page_number(fp)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)


def add_data_table(doc, headers, rows, widths, header_fill=PALE_JADE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, 10.5, bold=True, color=JADE)
    for row_values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row_values):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            r = p.add_run(text)
            set_run_font(r, 10.2, bold=(i == 0), color=CRIMSON if i == 0 else INK)
    set_table_geometry(table, widths)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    add_footer_and_header(section)
    first_footer = section.first_page_footer
    add_page_number(first_footer.paragraphs[0])

    flow_num = add_custom_numbering(doc, "decimal", "%1.")
    attr_bullets = add_custom_numbering(doc, "bullet", "•")
    battle_num = add_custom_numbering(doc, "decimal", "%1.")
    school_bullets = add_custom_numbering(doc, "bullet", "•")
    growth_bullets = add_custom_numbering(doc, "bullet", "•")
    tips_num = add_custom_numbering(doc, "decimal", "%1.")

    # Cover: editorial_cover pattern, compact-reference typography with a game-palette override.
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(6)
    kr = kicker.add_run("单人文学策略棋局  ·  玩家介绍")
    set_run_font(kr, 10, bold=True, color=GOLD)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("文 心 棋")
    set_run_font(tr, 30, bold=True, color=INK)
    remove_paragraph_borders(title)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("在桃花岛上走过三重科场，把一生文章写成自己的路")
    set_run_font(sr, 13, color=JADE)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(12)
    pic_run = pic_p.add_run()
    shape = pic_run.add_picture(str(COVER_ART), width=Inches(4.65))
    set_picture_alt(shape, "桃花岛书院", "桃花环绕的山水书院与石桥，象征玩家将进入的科场世界。")

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.left_indent = Inches(0.55)
    lead.paragraph_format.right_indent = Inches(0.55)
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.35
    lr = lead.add_run("掷骰行棋、答题取舍、挥毫论道。每一格都是际遇，每一次落笔都在塑造你的文名。")
    set_run_font(lr, 12, italic=True, color=CRIMSON)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    mr = meta.add_run("轻度 Roguelike  ·  棋盘行进  ·  策略构筑  ·  图鉴收集")
    set_run_font(mr, 9.5, bold=True, color=MUTED)

    doc.add_page_break()

    add_section_title(doc, "先读这一页：你会在游戏里做什么", "WELCOME TO WENXIN QI")
    add_callout(
        doc,
        "一句话认识它：",
        "《文心棋》是一款把大富翁式棋盘、文学创作对决与轻度 Roguelike 构筑揉在一起的单人游戏。你不是背完诗词就能赢；更重要的是判断时机、管理灵感，并逐步形成自己的写作流派。",
        fill=PALE_GOLD,
        border=GOLD,
    )

    doc.add_paragraph("一段不剧透的故事", style="Heading 2")
    add_body(
        doc,
        "你曾来自模糊的“现代世界”，在寻找世外仙源时抵达桃花岛。岛上仙人轻点灵台，你便成了蒙学馆中的小童生。十年寒窗之后，科举将启；你要从童生铺出发，走过层层试场，在一次次选择与论战中验证自己的文章，也追问那句旧日谶语：功名之外，究竟什么才算真正走出桃源？",
        keep=True,
    )

    doc.add_paragraph("一局游戏的完整节奏", style="Heading 2")
    add_numbered_step(doc, flow_num, "选流派。", "从博闻、奇士、辞宗中选择起点，决定开局侧重。")
    add_numbered_step(doc, flow_num, "装配名篇。", "从已解锁的传世名篇中携带至多 2 篇，让长期收藏转化为开局优势。")
    add_numbered_step(doc, flow_num, "掷骰行进。", "在三重同心棋盘上前进，落格触发恢复、成长、答题、奇遇、天象、名胜或论战。")
    add_numbered_step(doc, flow_num, "边走边构筑。", "提升六维、获取文心、积累心得与稿页，并决定灵感要花在何处。")
    add_numbered_step(doc, flow_num, "通过晋阶试。", "在阶段门迎战更强对手，从童生一路迈向秀才、举人与进士。")
    add_numbered_step(doc, flow_num, "参加殿试。", "穿过三圈后进入最终大考；满足特殊条件时，金榜之外还会出现一份隐藏终卷。")

    add_callout(
        doc,
        "游戏的核心乐趣：",
        "同一次掷骰会把你带到不同的资源与风险中；真正决定成败的，是你如何把随机结果变成可控的成长路线。",
        fill=PALE_JADE,
        border=JADE,
    )

    doc.add_page_break()

    add_section_title(doc, "棋盘、资源与六维能力", "THE BOARD")
    add_body(doc, "当前版本的三圈合计 192 格：外圈“起势”较宽松，中圈“验收”开始收紧，内圈“定稿”要求构筑成形。不同格子让每局路线与节奏都不一样。")

    add_data_table(
        doc,
        ["格子", "通俗作用"],
        [
            ("平韵", "恢复灵感，让你继续行路与创作。"),
            ("仄韵", "补充灵感，并逐步增强笔力、学力、思力。"),
            ("考题", "知识题考积累；创作抉择没有唯一正确答案，会留下你的“行卷”倾向。"),
            ("奇遇", "在收益、代价与剧情选择之间取舍，也可能获得稀有资源。"),
            ("论战", "进入“挥毫论道”，与不同位阶、不同弱点的文士比拼作品。"),
            ("天象", "改变接下来数回合的环境，例如强化答题、加重胜负或压制恢复。"),
            ("名胜", "可花费灵感，从 3 枚候选文心中挑选 1 枚；也可以保守离开。"),
        ],
        [1500, 7860],
    )

    doc.add_paragraph("灵感：既是体力，也是货币", style="Heading 2")
    add_callout(
        doc,
        "别把灵感只当生命值。",
        "你通常以 48 点灵感开局，上限 68。论战、奇遇、抽取文心和追加灵感骰都会消耗它；平韵、答题、技能与部分事件又能补回。灵感归零时，你将失去许多主动选择。",
        fill=PALE_RED,
        border=CRIMSON,
    )

    doc.add_paragraph("六维能力怎么理解", style="Heading 2")
    add_bullet(doc, attr_bullets, "诗力、词力、联力：三种创作文体的根基。结算既看三体共同底子，也看本体专精。", "诗力、词力、联力：")
    add_bullet(doc, attr_bullets, "笔力：让作品更有成色，也扩展稿匣、加快残稿变成稿页。", "笔力：")
    add_bullet(doc, attr_bullets, "学力：支撑意象与典故，同时提高心得容量和研修位。", "学力：")
    add_bullet(doc, attr_bullets, "思力：增强立意，并提供可提前设定的“章法”与构思次数。", "思力：")
    add_body(doc, "因此，你可以把一局玩成偏科的“一门宗师”，也可以追求三体与三功都稳定的全能路线。", bold_lead="因此，", after=0)

    doc.add_page_break()

    add_section_title(doc, "挥毫论道：一场文学对决怎么玩", "BATTLE IN SIX STEPS")
    add_body(doc, "对手并非单纯的数值墙。多数对手会公开偏好、意图或招牌能力，并留下可以利用的破绽；读懂提示，再决定文体与风格，往往比盲目堆分更重要。")

    add_numbered_step(doc, battle_num, "遭遇。", "先看对手的身份、偏科与研判提示。")
    add_numbered_step(doc, battle_num, "审题。", "确认题目、题材与本局“当朝风潮”；热点题材和得势风格可以叠加。")
    add_numbered_step(doc, battle_num, "选文体。", "在诗、词、联中选一种；此时要兼顾自身属性、上一场选择和对手弱点。")
    add_numbered_step(doc, battle_num, "选风格。", "清雅、豪放、婉约、沉郁等文风与题材存在相性，连续使用得当还能形成气势。")
    add_numbered_step(doc, battle_num, "掷灵感骰。", "先掷 1 枚；通常可再花 5 点灵感追加 1 枚，最多追加 2 枚，并得到额外作品增幅。")
    add_numbered_step(doc, battle_num, "算分对决。", "系统逐项展示格律、意象、立意、临场发挥与各类修正，让你看懂为什么赢、为什么输。")

    doc.add_paragraph("诗、词、联各有手感", style="Heading 2")
    add_callout(doc, "诗 · 一气：", "单骰高低分化明显，爆发高、波动也大；适合敢赌一笔的路线。", fill=PALE_RED, border=CRIMSON)
    add_callout(doc, "词 · 铺陈：", "首骰会收束到较稳定的区间，首次追加更省灵感；适合稳健运营。", fill=PALE_GOLD, border=GOLD)
    add_callout(doc, "联 · 对举：", "与上一场换文体时更强，失利也更能止损；适合观察局势、灵活换策。", fill=PALE_JADE, border=JADE)

    add_body(doc, "胜利当然重要，但失败也会获得熟练与心得。首局失利是允许的结果：这是一款鼓励你逐步认识系统、把下一局走得更远的游戏。", bold_lead="胜利当然重要，", after=0)

    doc.add_page_break()

    add_section_title(doc, "三大流派与构筑方向", "CHOOSE YOUR SCHOOL")
    add_body(doc, "流派决定你的开局侧重，却不会把路线锁死。你仍可以在棋盘上发展任何文体，或通过文心与名篇转向新的组合。")

    doc.add_paragraph("博闻｜博观约取，厚积薄发", style="Heading 2")
    add_callout(
        doc,
        "适合：",
        "喜欢答题、学习与均衡成长的玩家。开局学力更高，研修位更多；知识积累到一定程度后，还能获得可自由分配的心得。",
        fill="EEF2F7",
        border="4A6F96",
    )

    doc.add_paragraph("奇士｜灵台澄澈，万象皆明", style="Heading 2")
    add_callout(
        doc,
        "适合：",
        "喜欢计划、改线与把随机变成机会的玩家。开局思力更高，每阶段拥有更多构思，第一次章法发动还不消耗构思。",
        fill="F0EBF4",
        border="745C87",
    )

    doc.add_paragraph("辞宗｜笔落惊风雨，文成绣山川", style="Heading 2")
    add_callout(
        doc,
        "适合：",
        "喜欢稿本运营、把每场成果沉淀下来的玩家。开局笔力更高、稿匣更大；每阶段首次不追加骰完成论战，还能额外获得稿页。",
        fill="F5EEE7",
        border="8B5E3C",
    )

    doc.add_paragraph("一局之内，你还会经营四条成长线", style="Heading 2")
    add_bullet(doc, school_bullets, "文心：主动或被动技能。主动文心常需消耗灵感，被动文心持续改变规则；当前版本共有 43 枚。", "文心：")
    add_bullet(doc, school_bullets, "心得与研修：胜、平、负都能积累体悟，再把心得投向想加强的文体。", "心得与研修：")
    add_bullet(doc, school_bullets, "构思与章法：可设置徐行拾句、留白养气、换韵生新等计划，在满足条件时自动发动。", "构思与章法：")
    add_bullet(doc, school_bullets, "稿本：把残稿做成稿页，再选择润色、刊行或定卷，换取临场优势或最终得分。", "稿本：")

    add_callout(doc, "构筑没有标准答案：", "有时最强的选择不是把最高属性继续堆高，而是补足资源循环，让“拿到收益—转化成长—赢下论战”连成一条路。", fill=PALE_JADE, border=JADE)

    doc.add_page_break()

    add_section_title(doc, "跨局成长、结算与新手建议", "BEYOND A SINGLE RUN")

    doc.add_paragraph("每一局都会留下东西", style="Heading 2")
    add_bullet(doc, growth_bullets, "传世名篇：12 篇；解锁后可升至 Lv4 并二选一路线，选择会跨局保存。", "传世名篇：")
    add_bullet(doc, growth_bullets, "图鉴阁：记录对手、文心、羁绊与天象；当前还包含 9 组文心羁绊。", "图鉴阁：")
    add_bullet(doc, growth_bullets, "流派造诣：每局按表现获得熟练度；更高造诣不仅是荣誉，也与深层挑战有关。", "流派造诣：")
    add_bullet(doc, growth_bullets, "传承火种：少数特殊文心可把部分六维化作下一局起点，让失败也成为积累。", "传承火种：")

    doc.add_paragraph("如何判定这一局玩得怎么样", style="Heading 2")
    add_body(doc, "结算会从文采、功力、战绩、奇遇、流派、圆满六项汇总总分，并给出从童生到文宗的段位评价。你可以专精一体、追求三绝，也可以靠稳健资源管理与殿试表现拿到高分。")
    add_callout(
        doc,
        "关于结局：",
        "通过殿试即可金榜题名；若收集、造诣与殿试表现都足够出色，还会收到额外挑战“桃源终卷”的邀请，而且不会失去已经获得的金榜。",
        fill=PALE_GOLD,
        border=GOLD,
    )

    doc.add_paragraph("第一次玩，记住这 6 件事", style="Heading 2")
    add_numbered_step(doc, tips_num, "先保灵感循环。", "不要在每次论战都追骰；给名胜、关键战与意外事件留余量。", after=2)
    add_numbered_step(doc, tips_num, "看提示再出手。", "对手意图、弱点、题材与风潮都是公开信息，善用它们能省下大量数值。", after=2)
    add_numbered_step(doc, tips_num, "先有主轴，再补短板。", "围绕流派优势建立循环，同时避免某个关键资源完全断档。", after=2)
    add_numbered_step(doc, tips_num, "失败也要拿成长。", "换用不同文体可以积累熟练与心得；不必把每一场都当成非赢不可。", after=2)
    add_numbered_step(doc, tips_num, "阶段门前做整理。", "分配心得、使用稿页、确认章法与主动文心，别把资源留到来不及花。", after=2)
    add_numbered_step(doc, tips_num, "放心中断。", "每回合自动存档，也支持手动档与存档码；手机端可拖动、缩放棋盘。", after=2)

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(7)
    closing.paragraph_format.space_after = Pt(0)
    closing.paragraph_format.keep_together = True
    set_paragraph_shading(closing, PAPER)
    cr = closing.add_run("功名是一纸，文章是一生。\n愿你在《文心棋》中，走出属于自己的桃源。")
    set_run_font(cr, 11.5, bold=True, color=CRIMSON)

    doc.core_properties.title = "《文心棋》玩家介绍"
    doc.core_properties.subject = "面向玩家的剧情与玩法介绍"
    doc.core_properties.author = "《文心棋》项目组"
    doc.core_properties.keywords = "文心棋, 玩家介绍, 文学创作, 科举, Roguelike"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
